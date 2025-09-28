"""
Enhanced Export System for Legal Research Engine

This module provides comprehensive export capabilities for legal research results,
supporting multiple formats including Word, PDF, JSON, and custom legal formats.
"""

import streamlit as st
import json
import io
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from dataclasses import dataclass
import base64

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx not installed. Word export will be limited.")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    st.warning("reportlab not installed. PDF export will be limited.")

@dataclass
class ExportData:
    """Structure for legal research export data"""
    title: str
    query: str
    response: str
    sources: List[str]
    citations: List[str]
    confidence_score: float
    analysis_steps: List[str]
    precedents: List[str]
    recommendations: List[str]
    metadata: Dict[str, Any]
    timestamp: datetime

class LegalExportManager:
    """Manages export functionality for legal research results"""
    
    def __init__(self):
        self.supported_formats = ["txt", "json", "html"]
        
        if DOCX_AVAILABLE:
            self.supported_formats.append("docx")
        
        if REPORTLAB_AVAILABLE:
            self.supported_formats.append("pdf")
    
    def export_research_results(
        self, 
        data: ExportData, 
        format_type: str,
        include_metadata: bool = True,
        legal_format: str = "standard"
    ) -> bytes:
        """
        Export legal research results in the specified format
        
        Args:
            data: Research data to export
            format_type: Export format (txt, json, html, docx, pdf)
            include_metadata: Whether to include metadata in export
            legal_format: Legal formatting style (standard, bluebook, apa)
        
        Returns:
            Exported data as bytes
        """
        
        if format_type not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format_type}")
        
        if format_type == "txt":
            return self._export_to_text(data, include_metadata)
        elif format_type == "json":
            return self._export_to_json(data, include_metadata)
        elif format_type == "html":
            return self._export_to_html(data, include_metadata, legal_format)
        elif format_type == "docx" and DOCX_AVAILABLE:
            return self._export_to_word(data, include_metadata, legal_format)
        elif format_type == "pdf" and REPORTLAB_AVAILABLE:
            return self._export_to_pdf(data, include_metadata, legal_format)
        else:
            raise ValueError(f"Format {format_type} not available")
    
    def _export_to_text(self, data: ExportData, include_metadata: bool) -> bytes:
        """Export to plain text format"""
        
        content = f"""LEGAL RESEARCH REPORT
{'=' * 50}

Title: {data.title}
Query: {data.query}
Generated: {data.timestamp.strftime('%Y-%m-%d %H:%M:%S')}

ANALYSIS RESULTS
{'-' * 20}

{data.response}

"""
        
        if data.sources:
            content += f"""SOURCES
{'-' * 10}

"""
            for i, source in enumerate(data.sources, 1):
                content += f"{i}. {source}\n"
            content += "\n"
        
        if data.citations:
            content += f"""LEGAL CITATIONS
{'-' * 15}

"""
            for citation in data.citations:
                content += f"• {citation}\n"
            content += "\n"
        
        if data.precedents:
            content += f"""RELEVANT PRECEDENTS
{'-' * 20}

"""
            for precedent in data.precedents:
                content += f"• {precedent}\n"
            content += "\n"
        
        if data.recommendations:
            content += f"""RECOMMENDATIONS
{'-' * 15}

"""
            for rec in data.recommendations:
                content += f"• {rec}\n"
            content += "\n"
        
        if include_metadata:
            content += f"""METADATA
{'-' * 10}

Confidence Score: {data.confidence_score:.2%}
Analysis Steps: {len(data.analysis_steps)}
Processing Method: {data.metadata.get('method', 'Unknown')}

ANALYSIS STEPS:
"""
            for i, step in enumerate(data.analysis_steps, 1):
                content += f"{i}. {step}\n"
        
        return content.encode('utf-8')
    
    def _export_to_json(self, data: ExportData, include_metadata: bool) -> bytes:
        """Export to JSON format"""
        
        export_dict = {
            "title": data.title,
            "query": data.query,
            "response": data.response,
            "sources": data.sources,
            "citations": data.citations,
            "precedents": data.precedents,
            "recommendations": data.recommendations,
            "confidence_score": data.confidence_score,
            "timestamp": data.timestamp.isoformat(),
            "format_version": "1.0"
        }
        
        if include_metadata:
            export_dict.update({
                "analysis_steps": data.analysis_steps,
                "metadata": data.metadata
            })
        
        return json.dumps(export_dict, indent=2, ensure_ascii=False).encode('utf-8')
    
    def _export_to_html(self, data: ExportData, include_metadata: bool, legal_format: str) -> bytes:
        """Export to HTML format"""
        
        html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{data.title}</title>
    <style>
        body {{
            font-family: 'Times New Roman', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #fff;
            color: #333;
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        .header h1 {{
            color: #2c3e50;
            margin: 0;
            font-size: 2.2em;
        }}
        .meta-info {{
            color: #666;
            font-style: italic;
            margin-top: 10px;
        }}
        .section {{
            margin: 25px 0;
        }}
        .section-title {{
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 5px;
            margin-bottom: 15px;
            font-size: 1.3em;
            font-weight: bold;
        }}
        .query-box {{
            background-color: #f8f9fa;
            padding: 15px;
            border-left: 4px solid #3498db;
            margin: 15px 0;
            font-style: italic;
        }}
        .response-content {{
            background-color: #ffffff;
            padding: 20px;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            line-height: 1.8;
        }}
        .citations {{
            background-color: #f8f9fa;
            padding: 15px;
            border-left: 4px solid #28a745;
        }}
        .citation-item {{
            margin: 8px 0;
            padding-left: 15px;
            position: relative;
        }}
        .citation-item::before {{
            content: "§";
            position: absolute;
            left: 0;
            color: #28a745;
            font-weight: bold;
        }}
        .confidence-score {{
            background: linear-gradient(90deg, #28a745, #3498db);
            color: white;
            padding: 10px;
            border-radius: 5px;
            text-align: center;
            font-weight: bold;
        }}
        .metadata {{
            background-color: #f8f9fa;
            padding: 15px;
            border: 1px solid #dee2e6;
            border-radius: 5px;
            font-size: 0.9em;
        }}
        ul {{
            padding-left: 20px;
        }}
        li {{
            margin: 5px 0;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #dee2e6;
            text-align: center;
            color: #666;
            font-size: 0.9em;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Legal Research Report</h1>
        <div class="meta-info">
            Generated on {data.timestamp.strftime('%B %d, %Y at %I:%M %p')}
        </div>
    </div>
    
    <div class="section">
        <div class="section-title">Research Query</div>
        <div class="query-box">{data.query}</div>
    </div>
    
    <div class="section">
        <div class="section-title">Analysis Results</div>
        <div class="response-content">{data.response.replace(chr(10), '<br>')}</div>
    </div>
"""
        
        if data.confidence_score:
            html_content += f"""
    <div class="section">
        <div class="section-title">Confidence Assessment</div>
        <div class="confidence-score">
            Confidence Score: {data.confidence_score:.1%}
        </div>
    </div>
"""
        
        if data.citations:
            html_content += f"""
    <div class="section">
        <div class="section-title">Legal Citations</div>
        <div class="citations">
"""
            for citation in data.citations:
                html_content += f'            <div class="citation-item">{citation}</div>\n'
            
            html_content += """        </div>
    </div>
"""
        
        if data.sources:
            html_content += f"""
    <div class="section">
        <div class="section-title">Sources</div>
        <ul>
"""
            for source in data.sources:
                html_content += f'            <li>{source}</li>\n'
            
            html_content += """        </ul>
    </div>
"""
        
        if data.precedents:
            html_content += f"""
    <div class="section">
        <div class="section-title">Relevant Precedents</div>
        <ul>
"""
            for precedent in data.precedents:
                html_content += f'            <li>{precedent}</li>\n'
            
            html_content += """        </ul>
    </div>
"""
        
        if data.recommendations:
            html_content += f"""
    <div class="section">
        <div class="section-title">Recommendations</div>
        <ul>
"""
            for rec in data.recommendations:
                html_content += f'            <li>{rec}</li>\n'
            
            html_content += """        </ul>
    </div>
"""
        
        if include_metadata and data.analysis_steps:
            html_content += f"""
    <div class="section">
        <div class="section-title">Analysis Process</div>
        <div class="metadata">
            <ol>
"""
            for step in data.analysis_steps:
                html_content += f'                <li>{step}</li>\n'
            
            html_content += """            </ol>
        </div>
    </div>
"""
        
        html_content += f"""
    <div class="footer">
        <p>Legal Research Engine | Powered by LangGraph & Local AI Models</p>
        <p>This report was generated automatically and should be reviewed by a qualified legal professional.</p>
    </div>
</body>
</html>"""
        
        return html_content.encode('utf-8')
    
    def _export_to_word(self, data: ExportData, include_metadata: bool, legal_format: str) -> bytes:
        """Export to Word document format"""
        
        if not DOCX_AVAILABLE:
            raise ValueError("python-docx not available")
        
        doc = DocxDocument()
        
        # Document title
        title = doc.add_heading(data.title, 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Generated: {data.timestamp.strftime('%B %d, %Y at %I:%M %p')}").italic = True
        meta_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_page_break()
        
        # Query section
        doc.add_heading('Research Query', level=1)
        query_para = doc.add_paragraph(data.query)
        query_para.style = 'Quote'
        
        # Response section
        doc.add_heading('Analysis Results', level=1)
        doc.add_paragraph(data.response)
        
        # Confidence score
        if data.confidence_score:
            doc.add_heading('Confidence Assessment', level=2)
            doc.add_paragraph(f'Confidence Score: {data.confidence_score:.1%}')
        
        # Citations
        if data.citations:
            doc.add_heading('Legal Citations', level=1)
            for citation in data.citations:
                p = doc.add_paragraph(citation, style='List Bullet')
        
        # Sources
        if data.sources:
            doc.add_heading('Sources', level=1)
            for source in data.sources:
                doc.add_paragraph(source, style='List Number')
        
        # Precedents
        if data.precedents:
            doc.add_heading('Relevant Precedents', level=1)
            for precedent in data.precedents:
                doc.add_paragraph(precedent, style='List Bullet')
        
        # Recommendations
        if data.recommendations:
            doc.add_heading('Recommendations', level=1)
            for rec in data.recommendations:
                doc.add_paragraph(rec, style='List Bullet')
        
        # Analysis steps (if metadata included)
        if include_metadata and data.analysis_steps:
            doc.add_heading('Analysis Process', level=1)
            for i, step in enumerate(data.analysis_steps, 1):
                doc.add_paragraph(f"{i}. {step}")
        
        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io.getvalue()
    
    def _export_to_pdf(self, data: ExportData, include_metadata: bool, legal_format: str) -> bytes:
        """Export to PDF format"""
        
        if not REPORTLAB_AVAILABLE:
            raise ValueError("reportlab not available")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor('#2c3e50')
        )
        
        # Title
        story.append(Paragraph(data.title, title_style))
        story.append(Paragraph(f"Generated: {data.timestamp.strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Query
        story.append(Paragraph("Research Query", heading_style))
        story.append(Paragraph(data.query, styles['Quote']))
        story.append(Spacer(1, 12))
        
        # Response
        story.append(Paragraph("Analysis Results", heading_style))
        story.append(Paragraph(data.response, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Confidence score
        if data.confidence_score:
            story.append(Paragraph("Confidence Assessment", heading_style))
            story.append(Paragraph(f"Confidence Score: {data.confidence_score:.1%}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Citations
        if data.citations:
            story.append(Paragraph("Legal Citations", heading_style))
            for citation in data.citations:
                story.append(Paragraph(f"• {citation}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Sources
        if data.sources:
            story.append(Paragraph("Sources", heading_style))
            for i, source in enumerate(data.sources, 1):
                story.append(Paragraph(f"{i}. {source}", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer.getvalue()

def create_export_interface(research_data: Dict[str, Any]) -> None:
    """Create export interface in Streamlit"""
    
    st.subheader("📥 Export Research Results")
    
    # Export configuration
    col1, col2, col3 = st.columns(3)
    
    with col1:
        export_format = st.selectbox(
            "Export Format",
            options=["txt", "json", "html", "docx", "pdf"],
            format_func=lambda x: {
                "txt": "📄 Plain Text",
                "json": "🔧 JSON Data",
                "html": "🌐 Web Page", 
                "docx": "📝 Word Document",
                "pdf": "📋 PDF Report"
            }[x]
        )
    
    with col2:
        include_metadata = st.checkbox("Include Metadata", value=True)
    
    with col3:
        legal_format = st.selectbox(
            "Legal Format",
            options=["standard", "bluebook", "apa"],
            format_func=lambda x: x.title()
        )
    
    # Generate filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"legal_research_{timestamp}.{export_format}"
    
    if st.button("📥 Generate Export", type="primary"):
        try:
            with st.spinner("Generating export..."):
                # Convert research data to ExportData
                export_data = ExportData(
                    title=research_data.get("title", "Legal Research Report"),
                    query=research_data.get("query", ""),
                    response=research_data.get("response", ""),
                    sources=research_data.get("sources", []),
                    citations=research_data.get("citations", []),
                    confidence_score=research_data.get("confidence_score", 0.0),
                    analysis_steps=research_data.get("completed_steps", []),
                    precedents=research_data.get("precedents", []),
                    recommendations=research_data.get("recommendations", []),
                    metadata=research_data.get("metadata", {}),
                    timestamp=datetime.now()
                )
                
                # Export data
                exporter = LegalExportManager()
                exported_data = exporter.export_research_results(
                    export_data, 
                    export_format,
                    include_metadata,
                    legal_format
                )
                
                # Create download button
                st.download_button(
                    label=f"💾 Download {export_format.upper()} File",
                    data=exported_data,
                    file_name=filename,
                    mime=get_mime_type(export_format)
                )
                
                st.success(f"✅ Export generated successfully! ({len(exported_data)} bytes)")
                
        except Exception as e:
            st.error(f"Export failed: {str(e)}")

def get_mime_type(format_type: str) -> str:
    """Get MIME type for format"""
    mime_types = {
        "txt": "text/plain",
        "json": "application/json",
        "html": "text/html",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf"
    }
    return mime_types.get(format_type, "application/octet-stream")

def create_batch_export_interface() -> None:
    """Create interface for batch export of multiple research results"""
    
    st.subheader("📦 Batch Export")
    
    if "research_history" in st.session_state:
        history = st.session_state.research_history
        
        if history:
            # Select items for batch export
            selected_items = st.multiselect(
                "Select research results to export:",
                options=list(range(len(history))),
                format_func=lambda i: f"{history[i].get('query', 'Query')} ({history[i].get('timestamp', 'Unknown time')})"
            )
            
            if selected_items and st.button("📦 Create Batch Export"):
                # Create combined export
                combined_data = {
                    "title": "Batch Legal Research Report",
                    "query": f"Batch export of {len(selected_items)} research results",
                    "response": "Combined research results from multiple queries",
                    "batch_results": [history[i] for i in selected_items],
                    "timestamp": datetime.now()
                }
                
                create_export_interface(combined_data)
        else:
            st.info("No research history available for batch export")
    else:
        st.info("No research history found")

# Export main functions
__all__ = [
    'LegalExportManager',
    'ExportData', 
    'create_export_interface',
    'create_batch_export_interface'
]